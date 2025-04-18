import torch
import pandas as pd
import re
from flask import Flask, request, jsonify
from transformers import AutoTokenizer, AutoModelForCausalLM
from sentence_transformers import SentenceTransformer, util
from docx import Document
import mammoth
import tempfile
import os

# Load models
tokenizer = AutoTokenizer.from_pretrained("EleutherAI/gpt-neo-1.3B")
model = AutoModelForCausalLM.from_pretrained("EleutherAI/gpt-neo-1.3B")
if tokenizer.pad_token is None:
    tokenizer.add_special_tokens({'pad_token': '[PAD]'})
    model.resize_token_embeddings(len(tokenizer))
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model = model.to(device)

embedder = SentenceTransformer('all-MiniLM-L6-v2')

# Load data
CLAUSE_CSV_PATH = "Verified_50_Loan_Clause_Examples_final_fixed.csv"
TEMPLATE_PATH = "loan_agreement.docx"

prompt_df = pd.read_csv(CLAUSE_CSV_PATH)
prompt_df["parameters"] = prompt_df["parameters"].fillna("").astype(str)
prompt_df["parameters"] = prompt_df["parameters"].apply(
    lambda x: ", ".join(sorted(set(p.strip() for p in x.split(",") if p.strip().lower() != "nan")))
)
instruction_embeddings = embedder.encode(prompt_df['instruction'].tolist(), convert_to_tensor=True)

def build_prompt(example_1, example_2, instruction):
    return (
        "You are a legal assistant specialized in drafting formal legal clauses.\n"
        f"Example 1:\nClause: {example_1}\nEndClause\n\n"
        f"Example 2:\nClause: {example_2}\nEndClause\n\n"
        f"Now, generate ONLY the legal clause for {instruction}, using formal legal language.\n"
        "Output only the text between the markers 'Clause:' and 'EndClause'.\n\nClause: "
    )

def generate_clause(prompt_text):
    input_ids = tokenizer.encode(prompt_text, return_tensors="pt").to(device)
    output_ids = model.generate(
        input_ids,
        max_length=300,
        temperature=0.35,
        top_k=50,
        top_p=0.85,
        repetition_penalty=1.2,
        do_sample=True,
        num_return_sequences=1,
        pad_token_id=tokenizer.pad_token_id
    )
    generated_text = tokenizer.decode(output_ids[0], skip_special_tokens=True)
    if "EndClause" in generated_text:
        return generated_text.split("Clause:")[1].split("EndClause")[0].strip()
    return generated_text.split("Clause:")[1].strip()

def fill_parameters_dynamic(clause_text, param_string, provided_values):
    placeholders = set(re.findall(r"{(.*?)}", clause_text))
    defined_params = [p.strip() for p in str(param_string).split(',') if p.strip()]
    combined_params = sorted(placeholders.union(set(defined_params)))
    param_values = {}
    for param in combined_params:
        param_values[param] = provided_values.get(param, f"[{param}]")
    for param, value in param_values.items():
        clause_text = clause_text.replace(f"{{{param}}}", value)
    return clause_text

def find_best_match_semantic(user_input):
    user_embedding = embedder.encode(user_input, convert_to_tensor=True)
    cosine_scores = util.pytorch_cos_sim(user_embedding, instruction_embeddings)[0]
    best_idx = torch.argmax(cosine_scores).item()
    best_score = cosine_scores[best_idx].item()
    if best_score > 0.4:
        return prompt_df.iloc[best_idx]
    return None

app = Flask(__name__)

@app.route("/", methods=["GET"])
def index():
    return "Nyaya Jyoti AI Clause Generator Backend (DOCX + HTML via mammoth) is Running."

@app.route("/generate", methods=["POST"])
def generate():
    try:
        data = request.get_json()
        if not os.path.exists(TEMPLATE_PATH):
            return {"error": "Template not found."}, 500

        doc = Document(TEMPLATE_PATH)

        for key, value in data.items():
            if key != "Special_Clauses":
                for para in doc.paragraphs:
                    if f"[{key}]" in para.text:
                        para.text = para.text.replace(f"[{key}]", value)

        special_input = data.get("Special_Clauses", "").strip()
        special_clause_text = ""
        if special_input:
            match_row = find_best_match_semantic(special_input)
            if match_row is not None:
                prompt = build_prompt(match_row['example_1'], match_row['example_2'], match_row['instruction'])
                raw_clause = generate_clause(prompt)
                print("🧠 GPT RAW CLAUSE:", raw_clause)
                final_clause = fill_parameters_dynamic(raw_clause, match_row.get("parameters", ""), data)
                print("📋 Final Substituted Clause:", final_clause)
                special_clause_text = final_clause
            else:
                print("⚠️ No match found for special clause:", special_input)

        for para in doc.paragraphs:
            if "[Special_Clauses]" in para.text:
                para.text = para.text.replace("[Special_Clauses]", special_clause_text or "N/A")
                break

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            with open(tmp.name, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html_string = result.value
                docx_file.seek(0)
                encoded_docx = docx_file.read()

        return jsonify({
            "docx_base64": encoded_docx.decode("latin1"),
            "html_preview": html_string
        })

    except Exception as e:
        return {"error": str(e)}, 500

if __name__ == "__main__":
    app.run(port=5000)
